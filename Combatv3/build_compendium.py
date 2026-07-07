#!/usr/bin/env python3
"""build_compendium.py — render the Compendium .docx from compendium_data.json,
a pure-Python replacement for build_compendium.js (no Node required).

Matches the authored doc's look: landscape, 0.5" margins, EB Garamond, tight
bordered tables, **bold** markup parsed into runs.

Usage:  python build_compendium.py compendium_data.json Compendium.docx
"""
import json, re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "EB Garamond"
BODY = 8.0          # pt
HEAD_FILL = "EFEFEF"

def _set_cell_border(cell, color="auto", sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)

def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def _rich(paragraph, text, bold=False, italic=False, size=BODY):
    """Parse **bold** markup into runs on an existing paragraph."""
    text = str(text)
    parts = [p for p in re.split(r"(\*\*[^*]+\*\*)", text) if p != ""]
    if not parts:
        parts = [text]
    for p in parts:
        m = re.match(r"^\*\*([^*]+)\*\*$", p)
        r = paragraph.add_run(m.group(1) if m else p)
        r.font.name = FONT; r.font.size = Pt(size)
        r.bold = True if m else bold
        r.italic = italic

def _cell_para(cell):
    para = cell.paragraphs[0]
    pf = para.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    return para


def _content_weights(headers, rows, floor=6, cap=60):
    w=[]
    for i in range(len(headers)):
        m=len(re.sub(r"\*\*","",str(headers[i])))
        for r in rows:
            if i<len(r) and r[i] is not None:
                m=max(m,len(re.sub(r"\*\*","",str(r[i]))))
        w.append(min(cap,max(floor,m)))
    return w

def _set_col_widths(t, weights, total_in=10.0):
    tblPr=t._tbl.tblPr
    lay=OxmlElement("w:tblLayout"); lay.set(qn("w:type"),"fixed"); tblPr.append(lay)
    tot=float(sum(weights)) or 1.0
    for i,col in enumerate(t.columns):
        wtw=Twips(int(total_in*1440*weights[i]/tot))
        for cell in col.cells:
            cell.width=wtw

def add_table(doc, headers, rows, total_in=10.0, gap=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        p = _cell_para(hdr[i])
        _rich(p, h, bold=True)
        _shade(hdr[i], HEAD_FILL); _set_cell_border(hdr[i])
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if i >= len(cells): break
            p = _cell_para(cells[i])
            _rich(p, "" if val is None else val)
            _set_cell_border(cells[i])
    _set_col_widths(t, _content_weights(headers, rows), total_in)
    if gap:
        g=doc.add_paragraph(); g.paragraph_format.space_after=Pt(2)
    return t

def _heading(doc, text, size, before, after):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.keep_with_next = True
    r = p.add_run(text); r.bold = True; r.font.name = FONT; r.font.size = Pt(size)
    return p

def h1(doc, t): return _heading(doc, t, 15, 6, 2)
def h2(doc, t): return _heading(doc, t, 11, 3, 1)


def add_tables_row(doc, specs, widths_in=None):
    """specs=[(headers,rows),...] laid side by side to save vertical space."""
    n=len(specs)
    outer=doc.add_table(rows=1, cols=n); outer.alignment=WD_TABLE_ALIGNMENT.CENTER
    tblPr=outer._tbl.tblPr; lay=OxmlElement("w:tblLayout"); lay.set(qn("w:type"),"fixed"); tblPr.append(lay)
    page_in=10.0; widths_in=widths_in or [page_in/n]*n
    for i,(headers,rows) in enumerate(specs):
        cell=outer.rows[0].cells[i]
        cell.width=Twips(int(widths_in[i]*1440))
        # drop default empty para, build inner table inside the cell
        cell._tc.remove(cell.paragraphs[0]._p)
        it=cell.add_table(rows=1, cols=len(headers)); it.alignment=WD_TABLE_ALIGNMENT.LEFT; it.autofit=False
        hdr=it.rows[0].cells
        for j,h in enumerate(headers):
            p=_cell_para(hdr[j]); _rich(p,h,bold=True); _shade(hdr[j],HEAD_FILL); _set_cell_border(hdr[j])
        for row in rows:
            cc=it.add_row().cells
            for j,val in enumerate(row):
                if j>=len(cc): break
                _rich(_cell_para(cc[j]),"" if val is None else val); _set_cell_border(cc[j])
        _set_col_widths(it,_content_weights(headers,rows),widths_in[i]-0.1)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return outer

def build(data, out_path):
    doc = Document()
    # default font
    style = doc.styles["Normal"]
    style.font.name = FONT; style.font.size = Pt(BODY)
    # landscape, 0.5" margins
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Twips(16838); sec.page_height = Twips(11906)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Twips(504))

    # Title
    p = doc.add_paragraph(); r = p.add_run("Renown — Compendium")
    r.bold = True; r.font.name = FONT; r.font.size = Pt(22)
    ver = data.get("version", "")
    sub = "Generated from renown_data.py — the single source of truth."
    if ver:
        sub = f"v{ver} · " + sub
    p2 = doc.add_paragraph(); _rich(p2, sub, italic=True)

    # Pursuits
    h1(doc, "Pursuits")
    for s in data["pursuit_sections"]:
        h2(doc, s["title"])
        add_table(doc, ["Pursuit", "Mastery Unlock", "Innate Effect", "Mastery Effect"], s["rows"])

    # Equipment
    h1(doc, "Equipment")
    eq = data["equipment"]
    h2(doc, "Retinues");       add_table(doc, ["Retinue","Cost","To Hit","Endurance","Morale","Keyword"], eq["Retinues"])
    h2(doc, "Melee Weapons");  add_table(doc, ["Weapon","Tier","AP","Init","Keywords"], eq["Weapons"])
    h2(doc, "Ranged Weapons"); add_table(doc, ["Ranged","Tier","AP","Init","Keywords"], eq["Ranged"])
    h2(doc, "Shields");        add_table(doc, ["Shield","Tier","Save","Init","Keywords"], eq["Shields"])
    h2(doc, "Armor");          add_table(doc, ["Armor","Tier","Save","Keywords"], eq["Armor"])

    # Infrastructure & Wonders
    h1(doc, "Infrastructure & Wonders")
    add_table(doc, ["Infrastructure","Upkeep","Freq","Empire Bonus","Tier","Build","Requirement"], data["infrastructure"])
    h2(doc, "Wonders"); add_table(doc, ["Wonder","Empire Bonus","Build","Requirement"], data["wonders"])

    # Empire
    h1(doc, "Empire")
    h2(doc, "Settlements"); add_table(doc, ["Settlement","Tier","Sea Variant","Tax","Muster","Build","Wards","Reach","Notes"], data["settlements"])
    h2(doc, "Eras");        add_table(doc, ["Era","Renown","Armies","Cities","Infl/Turn","Diplo Infl","Envoys","Unlocks"], data["eras"])
    h2(doc, "Domain Standings — empire + combat")
    _emp = data["domain_board"]
    _cmb = {r[0]: r for r in data["standing_effects"]}   # keyed by domain
    _merged = []
    for er in _emp:
        dom = er[0]; cr = _cmb.get(dom, [dom, "", "", ""])
        row = [dom]
        for i in (1, 2, 3):
            emp = (er[i] if i < len(er) and er[i] else "").strip()
            cmb = (cr[i] if i < len(cr) and cr[i] else "").strip()
            cell = emp
            if cmb:
                cell = (emp + "  " if emp else "") + f"**Combat:** {cmb}"
            row.append(cell)
        _merged.append(row)
    add_table(doc, ["Domain", "Rising (3)", "Established (6)", "Sovereign (10)"], _merged)
    h2(doc, "Tactic Matrix"); add_table(doc, data["tactic_matrix_header"], data["tactic_matrix_rows"])
    h2(doc, "Public Order"); add_table(doc, ["PO","State","Effect"], data["public_order"])
    h2(doc, "Faith & Doubt · Seasons · Trade")
    add_tables_row(doc, [(["Type","Source","Condition"], data["po_modifiers"]),
                         (["Season","Name","Effect"], data["seasons"]),
                         (["Rule","Value"], data["trade_rules"])],
                   widths_in=[4.0,3.2,2.8])

    # Factions
    h1(doc, "Factions")
    add_table(doc, ["Faction","Mechanic"], data["factions"])

    # Glossary
    h1(doc, "Glossary")
    for cat in data["glossary_categorized"]:
        h2(doc, cat["title"])
        add_table(doc, ["Term","Definition"], cat["rows"])

    doc.save(out_path)
    print("wrote " + out_path)

if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else "compendium_data.json"
    out = sys.argv[2] if len(sys.argv) > 2 else "Compendium.docx"
    build(json.load(open(src, encoding="utf-8")), out)