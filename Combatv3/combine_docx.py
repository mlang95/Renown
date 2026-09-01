#!/usr/bin/env python3
"""combine_docx.py — append the Compendium after the Rules doc as one book,
preserving each part's own section (portrait rules -> landscape reference), and
stamp a footer on every section: "Renown v{VERSION}  .  Page X of Y".

Version is read from renown_data.VERSION; pass a 4th arg to override.

Inputs are produced elsewhere in :docs
  md_to_docx.py RULES_reorganized_5.md Rules.docx
  build_compendium.py compendium_data.json Compendium.docx

Usage:  python combine_docx.py Rules.docx Compendium.docx Renown.docx [version]
"""
import sys
from copy import deepcopy
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

FONT = "EB Garamond"
SIZE = 8  # pt

try:
    import renown_data as _rd
    VERSION = str(getattr(_rd, "VERSION", ""))
except Exception:
    VERSION = ""


def _rpr(font=FONT, size=SIZE):
    rpr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), font)
    rpr.append(rf)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rpr.append(sz)
    return rpr


def _field(instr):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    r = OxmlElement("w:r")
    r.append(_rpr())
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    return fld


def _text(text):
    r = OxmlElement("w:r")
    r.append(_rpr())
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def combine(rules_path, comp_path):
    """Merge into one Document (portrait rules -> landscape reference)."""
    A = Document(rules_path)
    B = Document(comp_path)
    ab, bb = A.element.body, B.element.body
    # Close A's (portrait) section by moving its body-level sectPr into a trailing paragraph...
    a_sectPr = ab.findall(qn("w:sectPr"))[-1]
    p = A.add_paragraph()
    p._p.get_or_add_pPr().append(deepcopy(a_sectPr))
    ab.remove(a_sectPr)
    # ...then append B's content, which ends in B's own (landscape) sectPr = final section.
    for child in list(bb):
        ab.append(deepcopy(child))
    return A


def add_footers(doc, version):
    for section in doc.sections:
        section.footer.is_linked_to_previous = False
        p = section.footer.paragraphs[0]
        for r in list(p._p.findall(qn("w:r"))):
            p._p.remove(r)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p._p.append(_text(f"Renown v{version}   \u00b7   Page "))
        p._p.append(_field("PAGE"))
        p._p.append(_text(" of "))
        p._p.append(_field("NUMPAGES"))


if __name__ == "__main__":
    if len(sys.argv) < 4:
        sys.exit("usage: python combine_docx.py Rules.docx Compendium.docx Renown.docx [version]")
    version = sys.argv[4] if len(sys.argv) > 4 else VERSION
    doc = combine(sys.argv[1], sys.argv[2])
    add_footers(doc, version)
    doc.save(sys.argv[3])
    print(f"wrote {sys.argv[3]} (footer: Renown v{version} + page X of Y)")