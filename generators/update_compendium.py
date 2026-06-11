"""
update_compendium.py — drives Compendium.docx from the canonical CSVs.

Usage (in a Jupyter cell or script):
    from update_compendium import update_compendium
    update_compendium(
        compendium_in   = r"C:\\path\\Compendium.docx",
        compendium_out  = r"C:\\path\\Compendium-Updated.docx",
        specs_csv       = "specs.csv",
        infrastructure_csv = "infrastructure.csv",
        equipment_csv   = "equipment.csv",
    )

What this does
--------------
Opens an existing Compendium.docx and updates the *cell text* of the
data tables (specs, infrastructure, wonders, equipment). It does NOT
re-create tables — column widths, row heights, padding, and any
VBA-tuned layout are preserved because we modify cells in place.

Prose paragraphs, the glossary, factions list, terrain, settlement
tier table, and other reference content are NOT touched.

The script reports any drift it finds:
- Specs in the CSV that don't exist in the docx
- Specs in the docx that aren't in the CSV
- Equipment/infrastructure items that don't match
"""

import csv
import re
from copy import deepcopy
from docx import Document
from docx.shared import Pt


# ----------------------------------------------------------------------
# Spec-name parsing
# ----------------------------------------------------------------------
# Docx cells use a composite format for spec names:
#   "Name"                                  — plain
#   "Fishmongery ^"                         — water-settlement marker
#   "Reliquary * (Rising Piety)"            — unique + tier
#   "Coliseum (Rising Prowess)"             — tier only
#   "Royal Pavilion * (Sovereign Prowess)"  — unique + tier
#   "Master Workshop*"                       — asterisk attached (typo)
#   "Shipwright *^"                          — combined markers
#   "Gilded Foundry * Established Industry)" — missing open paren (typo)
#
# We parse these so we can match against the canonical Specialization
# field in specs.csv, which has none of these markers.

# Aliases: docx name → CSV name (for naming inconsistencies)
NAME_ALIASES = {
    "School of Engineering": "College of Engineering",
}


def parse_spec_cell(cell_text):
    """Parse the composite spec cell. Returns dict with name and markers.
    Robust against docx typos like 'Master Workshop*' or missing open paren."""
    s = cell_text.strip()
    if not s:
        return {"name": "", "unique": False, "water": False, "tier_label": ""}

    # Detect & strip a trailing parenthetical tier label.
    # Accept both "(Rising Piety)" and the typo form "Established Industry)"
    tier_label = ""
    m = re.search(r"\(([^)]*)\)\s*$", s)
    if m:
        tier_label = m.group(1).strip()
        s = s[:m.start()].rstrip()
    else:
        # Typo: closing paren without opening one. Pull everything from the
        # last asterisk-space onward as a candidate tier label.
        m2 = re.search(r"\*\s+([A-Z][A-Za-z\s]+)\)\s*$", s)
        if m2:
            tier_label = m2.group(1).strip()
            s = s[:m2.start() + 1].rstrip()  # keep the *

    # Now strip markers (* and ^), wherever they sit
    unique = "*" in s
    water = "^" in s
    name = s.replace("*", "").replace("^", "").strip()
    # Collapse double-space
    name = re.sub(r"\s+", " ", name)
    # Resolve aliases
    name = NAME_ALIASES.get(name, name)

    return {"name": name, "unique": unique, "water": water, "tier_label": tier_label}


def format_spec_cell(name, unique=False, water=False, tier_label=""):
    """Reconstruct the docx cell label from components."""
    s = name
    if water:
        s += " ^"
    if unique:
        s += " *"
    if tier_label:
        s += f" ({tier_label})"
    return s


def derive_tier_label(unlock_requirement):
    """From specs.csv 'Unlock Requirement' column, build the parenthetical
    used in the docx, or return '' if it shouldn't have one."""
    u = (unlock_requirement or "").strip()
    if not u or u == "-":
        return ""
    # Skip non-domain unlocks like "Water Settlement"
    if u.lower() == "water settlement":
        return ""
    # Composite unlocks like "Sovereign Piety + Established Prowess"
    # or "Established Industry, Water Settlement"
    return u


def is_unique_type(spec_type):
    """Per Compendium convention, Power and Monument and certain capstones
    are unique. We'll mark these with *."""
    return spec_type in ("Power", "Monument")


def is_water_spec(unlock_requirement):
    return "Water Settlement" in (unlock_requirement or "")


# ----------------------------------------------------------------------
# Cell writing — preserves formatting
# ----------------------------------------------------------------------
def _clear_paragraph(paragraph):
    """Remove all runs from a paragraph (but keep the paragraph and its props)."""
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def _set_cell_text(cell, text, template_run=None):
    """Replace cell text. Recognizes **bold** markers in text and produces
    matching bold runs. Preserves the cell's first paragraph's properties
    (alignment, indentation) and copies formatting from template_run if given,
    else from the cell's existing first run (font name, size).
    """
    if text is None:
        text = ""

    # Find or capture a template run for formatting
    if template_run is None:
        for p in cell.paragraphs:
            for r in p.runs:
                if r.text.strip():
                    template_run = r
                    break
            if template_run is not None:
                break

    # Keep the first paragraph, drop the rest, clear its runs
    paragraphs = list(cell.paragraphs)
    main_p = paragraphs[0]
    for extra_p in paragraphs[1:]:
        extra_p._element.getparent().remove(extra_p._element)
    _clear_paragraph(main_p)

    # Tokenize on **bold** markers
    tokens = re.split(r'(\*\*[^*]+?\*\*)', text)
    for tok in tokens:
        if not tok:
            continue
        is_bold = tok.startswith('**') and tok.endswith('**')
        inner = tok[2:-2] if is_bold else tok
        if not inner:
            continue
        run = main_p.add_run(inner)
        if template_run is not None:
            # Copy template's font properties
            if template_run.font.name:
                run.font.name = template_run.font.name
            if template_run.font.size:
                run.font.size = template_run.font.size
        run.bold = True if is_bold else None


# ----------------------------------------------------------------------
# Auto-detect table indices by header content
# ----------------------------------------------------------------------
# We no longer hardcode which docx table index corresponds to which data
# section, because the docx structure shifts when types get added/removed.
# Instead we walk the tables and identify each one by its header row and
# (for spec tables) by sampling the first data row's spec type.


def _build_table_map(doc, specs_by_name, eq_by_cat_name=None):
    """Walk doc.tables and identify each table's purpose.
    Returns dict mapping table index -> {"kind": str, "type": Optional[str]}.

    Spec tables are detected by the 4-column header 'Specialization /
    Mastery Unlock / Innate Effect / Mastery Effect', and their CSV type is
    inferred by looking up the first non-header data row's spec name in
    specs_by_name.

    Weapons and Ranged both use 'Weapon' as their first header, so they
    are distinguished by inspecting the first data row's name against
    equipment.csv categories (if provided).
    """
    table_map = {}
    eq_by_cat_name = eq_by_cat_name or {}
    ranged_names = set(eq_by_cat_name.get("Ranged", {}).keys())
    weapon_names = set(eq_by_cat_name.get("Weapon", {}).keys())

    for i, t in enumerate(doc.tables):
        if not t.rows or not t.rows[0].cells:
            continue
        header = [c.text.strip() for c in t.rows[0].cells]
        first_data_name = ""
        if len(t.rows) > 1 and t.rows[1].cells:
            first_data_name = t.rows[1].cells[0].text.strip().replace("\u2019", "'")

        # Spec tables: 4 columns starting with 'Specialization' or 'Pursuit'
        if len(header) >= 4 and header[0].lower() in ('specialization', 'pursuit') and \
           'mastery unlock' in header[1].lower():
            spec_type = None
            for row in t.rows[1:]:
                if not row.cells:
                    continue
                parsed = parse_spec_cell(row.cells[0].text)
                spec = specs_by_name.get(parsed["name"])
                if spec:
                    spec_type = spec["Type"]
                    break
            table_map[i] = {"kind": "spec", "type": spec_type}
        elif header[0].lower() == 'infrastructure':
            table_map[i] = {"kind": "infrastructure", "type": None}
        elif header[0].lower() == 'wonder':
            table_map[i] = {"kind": "wonder", "type": None}
        elif header[0].lower() == 'weapon':
            # Distinguish weapons (melee) vs ranged using first data row name
            if first_data_name in ranged_names:
                table_map[i] = {"kind": "ranged", "type": None}
            elif first_data_name in weapon_names:
                table_map[i] = {"kind": "weapon", "type": None}
            else:
                # No equipment.csv passed yet — fall back to header column count
                # Both have 6 columns, so default to weapon if ambiguous;
                # this branch is only hit during the bootstrap call.
                table_map[i] = {"kind": "weapon", "type": None}
        elif header[0].lower() == 'shield':
            table_map[i] = {"kind": "shield", "type": None}
        elif header[0].lower() == 'armor':
            table_map[i] = {"kind": "armor", "type": None}
        elif 'retinue' in header[0].lower():
            table_map[i] = {"kind": "retinue", "type": None}
        elif header[0].lower() == 'tier' and len(header) > 2 and 'specialization' in header[1].lower():
            table_map[i] = {"kind": "tier_requirements", "type": None}
    return table_map


# ----------------------------------------------------------------------
# Spec table updates
# ----------------------------------------------------------------------


def update_spec_tables(doc, specs_by_name, table_map):
    """For each spec table identified in table_map, walk its data rows and
    update Mastery Unlock / Innate / Mastery from the CSV.
    Reports any mismatches found.
    Returns lists: (updated, missing_in_csv, missing_in_docx)."""
    updated = []
    missing_in_csv = []
    docx_names_seen = set()

    for table_idx, info in table_map.items():
        if info["kind"] != "spec":
            continue
        table = doc.tables[table_idx]
        expected_type = info["type"]
        for row in table.rows[1:]:  # skip header
            cells = row.cells
            if len(cells) < 4:
                continue
            parsed = parse_spec_cell(cells[0].text)
            name = parsed["name"]
            if not name:
                continue
            docx_names_seen.add(name)

            spec = specs_by_name.get(name)
            if spec is None:
                missing_in_csv.append(
                    f"Table {table_idx} ({expected_type or '?'}): docx row '{name}' has no matching CSV spec"
                )
                continue

            type_warning = None
            if expected_type and spec["Type"] != expected_type:
                type_warning = (
                    f"Table {table_idx} ({expected_type}): "
                    f"docx row '{name}' has CSV type '{spec['Type']}' — "
                    f"consider moving this spec to the {spec['Type']} table"
                )

            _set_cell_text(cells[1], spec["Mastery Requirement"])
            _set_cell_text(cells[2], spec["Innate Effects"])
            _set_cell_text(cells[3], spec["Mastery Effect"])

            new_label = format_spec_cell(
                name,
                unique=is_unique_type(spec["Type"]),
                water=is_water_spec(spec["Unlock Requirement"]),
                tier_label=derive_tier_label(spec["Unlock Requirement"]),
            )
            _set_cell_text(cells[0], new_label)

            if type_warning:
                missing_in_csv.append(type_warning)
            updated.append(name)

    missing_in_docx = []
    for spec_name, spec in specs_by_name.items():
        if spec_name in docx_names_seen:
            continue
        missing_in_docx.append(
            f"CSV spec '{spec_name}' ({spec['Type']}) is not in any docx spec table"
        )

    return updated, missing_in_csv, missing_in_docx


# ----------------------------------------------------------------------
# Infrastructure & Wonder table updates
# ----------------------------------------------------------------------
def update_infrastructure_table(doc, infra_rows, table_map):
    """Update the docx Infrastructure table from infrastructure.csv where Category=Infrastructure.
    Column layout: Infrastructure | Upkeep | Upkeep Frequency | Empire Bonus | Tier | Build Time | Requirement
    """
    issues = []
    infra_by_name = {r["Name"]: r for r in infra_rows if r["Category"] == "Infrastructure"}

    infra_table_idx = None
    for idx, info in table_map.items():
        if info["kind"] == "infrastructure":
            infra_table_idx = idx
            break
    if infra_table_idx is None:
        issues.append("No Infrastructure table found in docx")
        return issues

    table = doc.tables[infra_table_idx]
    seen = set()
    for row in table.rows[1:]:
        cells = row.cells
        name = cells[0].text.strip()
        # Skip the generic 'Wonder' placeholder row — Wonders are tracked in their own table.
        if name == "Wonder":
            continue
        seen.add(name)
        if name not in infra_by_name:
            issues.append(f"Table {infra_table_idx}: docx row '{name}' not in infrastructure.csv")
            continue
        r = infra_by_name[name]
        _set_cell_text(cells[1], r["Upkeep"])
        _set_cell_text(cells[2], r["Upkeep Frequency"])
        _set_cell_text(cells[3], r["Empire Bonus"])
        _set_cell_text(cells[4], r["Tier"])
        _set_cell_text(cells[5], r["Build Time"])
        _set_cell_text(cells[6], r["Requirement"])

    for name in infra_by_name:
        if name not in seen:
            issues.append(f"Infrastructure.csv row '{name}' not in docx Infrastructure table")
    return issues


def update_wonders_table(doc, infra_rows, table_map):
    """Update the docx Wonders table from infrastructure.csv where Category=Wonder."""
    issues = []
    wonder_by_name = {r["Name"]: r for r in infra_rows if r["Category"] == "Wonder"}

    wonder_table_idx = None
    for idx, info in table_map.items():
        if info["kind"] == "wonder":
            wonder_table_idx = idx
            break
    if wonder_table_idx is None:
        issues.append("No Wonders table found in docx")
        return issues

    table = doc.tables[wonder_table_idx]
    seen = set()
    for row in table.rows[1:]:
        cells = row.cells
        name = cells[0].text.strip()
        seen.add(name)
        if name not in wonder_by_name:
            issues.append(f"Table {wonder_table_idx}: docx wonder '{name}' not in infrastructure.csv")
            continue
        r = wonder_by_name[name]
        _set_cell_text(cells[1], r["Empire Bonus"])

    for name in wonder_by_name:
        if name not in seen:
            issues.append(f"Infrastructure.csv wonder '{name}' not in docx Wonders table")
    return issues


# ----------------------------------------------------------------------
# Equipment table updates
# ----------------------------------------------------------------------
EQUIPMENT_COLS = {
    "weapon":  ["AP", "Initiative", "Effects", "Tier", "Specialization Unlock"],
    "ranged":  ["AP", "Initiative", "Effects", "Tier", "Specialization Unlock"],
    "shield":  ["Save", "Initiative", "Effects", "Tier", "Specialization Unlock", "Note"],
    "armor":   ["Save", "Tier"],
    "retinue": ["Cost", "To Hit", "Endurance", "Shaking", "Specialization Unlock"],
}

EQUIPMENT_CSV_CATEGORIES = {
    "weapon": "Weapon",
    "ranged": "Ranged",
    "shield": "Shield",
    "armor": "Armor",
    "retinue": "Retinue",
}


def update_equipment_tables(doc, equipment_rows, table_map):
    """Update the docx equipment tables from equipment.csv.
    Auto-detected by kind: weapon, ranged, shield, armor, retinue."""
    issues = []
    eq_by_cat_name = {}
    for r in equipment_rows:
        eq_by_cat_name.setdefault(r["Category"], {})[r["Name"]] = r

    for table_idx, info in table_map.items():
        kind = info["kind"]
        if kind not in EQUIPMENT_COLS:
            continue
        col_order = EQUIPMENT_COLS[kind]
        csv_cat = EQUIPMENT_CSV_CATEGORIES[kind]
        table = doc.tables[table_idx]
        items = eq_by_cat_name.get(csv_cat, {})
        seen = set()

        for row in table.rows[1:]:
            cells = row.cells
            name = cells[0].text.strip()
            name = name.replace("\u2019", "'")
            seen.add(name)
            if name not in items:
                issues.append(f"Table {table_idx} ({csv_cat}): docx row '{name}' not in equipment.csv")
                continue
            r = items[name]
            for i, col_name in enumerate(col_order):
                if i + 1 >= len(cells):
                    break
                val = r.get(col_name, "") or ""
                _set_cell_text(cells[i + 1], val)
            _set_cell_text(cells[0], r["Name"])

        for name in items:
            if name not in seen:
                issues.append(f"Equipment.csv {csv_cat} '{name}' not in docx Table {table_idx}")
    return issues


# ----------------------------------------------------------------------
# Main
# ----------------------------------------------------------------------
def load_csv(path):
    with open(path, newline="", encoding="cp1252") as f:
        return list(csv.DictReader(f))


def update_compendium(compendium_in, compendium_out,
                      specs_csv, infrastructure_csv, equipment_csv,
                      verbose=True):
    specs_rows = load_csv(specs_csv)
    infra_rows = load_csv(infrastructure_csv)
    eq_rows = load_csv(equipment_csv)

    specs_by_name = {r["Pursuits"]: r for r in specs_rows}

    # Build equipment lookup for weapon/ranged disambiguation
    eq_by_cat_name = {}
    for r in eq_rows:
        eq_by_cat_name.setdefault(r["Category"], {})[r["Name"]] = r

    doc = Document(compendium_in)

    # Detect table structure dynamically so the script survives docx reorgs.
    table_map = _build_table_map(doc, specs_by_name, eq_by_cat_name)

    updated, missing_in_csv, missing_in_docx = update_spec_tables(doc, specs_by_name, table_map)
    infra_issues = update_infrastructure_table(doc, infra_rows, table_map)
    wonder_issues = update_wonders_table(doc, infra_rows, table_map)
    equip_issues = update_equipment_tables(doc, eq_rows, table_map)

    doc.save(compendium_out)

    if verbose:
        print(f"Compendium written to: {compendium_out}")
        print(f"Specs updated: {len(updated)}")
        # Show the auto-detected table map for transparency
        spec_tables = [(i, info["type"]) for i, info in table_map.items() if info["kind"] == "spec"]
        if spec_tables:
            print(f"Detected spec tables: " + ", ".join(f"#{i}={t}" for i, t in spec_tables))
        if missing_in_csv:
            print(f"\n{len(missing_in_csv)} docx specs not matched in specs.csv:")
            for m in missing_in_csv:
                print(f"  · {m}")
        if missing_in_docx:
            print(f"\n{len(missing_in_docx)} CSV specs not in docx tables:")
            for m in missing_in_docx:
                print(f"  · {m}")
        if infra_issues:
            print(f"\n{len(infra_issues)} Infrastructure issues:")
            for m in infra_issues:
                print(f"  · {m}")
        if wonder_issues:
            print(f"\n{len(wonder_issues)} Wonder issues:")
            for m in wonder_issues:
                print(f"  · {m}")
        if equip_issues:
            print(f"\n{len(equip_issues)} Equipment issues:")
            for m in equip_issues:
                print(f"  · {m}")
        if not (missing_in_csv or missing_in_docx or infra_issues or wonder_issues or equip_issues):
            print("All tables aligned.")


if __name__ == "__main__":
    update_compendium(
        compendium_in="Compendium.docx",
        compendium_out="Compendium-Updated.docx",
        specs_csv="specs.csv",
        infrastructure_csv="infrastructure.csv",
        equipment_csv="equipment.csv",
    )
