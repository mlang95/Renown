#!/usr/bin/env python3
"""combine_docx.py — append the Compendium after the Rules doc as one book:
  - portrait rules  ->  landscape reference (each keeps its own section)
  - a master Table of Contents at the front (whole book; from outline levels)
  - a second Table of Contents at the start of the compendium (compendium only;
    built from hidden TC entry-marks flagged \\f C, which Word AND LibreOffice honor)
  - a footer on every section: "Renown v{VERSION}  .  Page X of Y"
  - flags fields to refresh when the document is opened

Both TOCs need the headings to carry w:outlineLvl (see the one-line edits to
md_to_docx.py and build_compendium.py shipped alongside this file). The master
TOC reads those levels directly; the compendium TOC reads the TC marks this
script injects into each compendium heading.

Version is read from renown_data.VERSION; pass a 5th arg to override.

Usage:  python combine_docx.py Rules.docx Compendium.docx Renown.docx [version]
"""
import sys
from copy import deepcopy
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

FONT = "EB Garamond"
SIZE = 8            # footer pt
COMP_FLAG = "C"     # TC \f flag identifying compendium entries

try:
    import renown_data as _rd
    VERSION = str(getattr(_rd, "VERSION", ""))
except Exception:
    VERSION = ""

W_T = qn("w:t"); W_P = qn("w:p"); W_PPR = qn("w:pPr"); W_OL = qn("w:outlineLvl"); W_SECT = qn("w:sectPr")


# ── run/field helpers ────────────────────────────────────────────────────────
def _rpr(font=FONT, size=SIZE, bold=False):
    rpr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), font)
    rpr.append(rf)
    if bold:
        rpr.append(OxmlElement("w:b"))
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rpr.append(sz)
    return rpr


def _run(text=None, fldchar=None, instr=None, font=FONT, size=SIZE, bold=False):
    r = OxmlElement("w:r")
    r.append(_rpr(font, size, bold))
    if fldchar is not None:
        fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), fldchar)
        if fldchar == "begin":
            fc.set(qn("w:dirty"), "true")
        r.append(fc)
    elif instr is not None:
        it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
        r.append(it)
    else:
        t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
        r.append(t)
    return r


def _para(*runs, after=None):
    p = OxmlElement("w:p")
    if after is not None:
        pPr = OxmlElement("w:pPr")
        sp = OxmlElement("w:spacing"); sp.set(qn("w:after"), str(after)); pPr.append(sp)
        p.append(pPr)
    for r in runs:
        p.append(r)
    return p


def _page_break():
    r = OxmlElement("w:r")
    br = OxmlElement("w:br"); br.set(qn("w:type"), "page"); r.append(br)
    return _para(r)


def _toc_block(title_text, levels="1-2", flag=None, title_size=18):
    if flag:
        instr = f' TOC \\f {flag} \\l "{levels}" \\h \\z '
    else:
        instr = f' TOC \\o "{levels}" \\h \\z \\u '
    title = _para(_run(title_text, size=title_size, bold=True), after=120)
    toc = _para(
        _run(fldchar="begin"),
        _run(instr=instr),
        _run(fldchar="separate"),
        _run("Right-click and Update Field (or press F9) to build the table of contents.", size=10),
        _run(fldchar="end"),
    )
    return [title, toc, _page_break()]


def _appendix_heading(text="Appendix \u2014 The Compendium"):
    """A level-1 divider the master TOC picks up as a single appendix entry."""
    p = _para(_run(text, size=18, bold=True), after=120)
    ol = OxmlElement("w:outlineLvl"); ol.set(qn("w:val"), "0")
    p.find(W_PPR).append(ol)
    return p


def _tag_heading_as_tc(p, flag):
    """Mark a compendium heading for the compendium-only TOC (TC \\f flag), and strip
    its outlineLvl so the master (\\o) TOC skips it — the compendium collapses to one
    'Appendix' entry up front while keeping full detail in its own TOC."""
    pPr = p.find(W_PPR)
    ol = pPr.find(W_OL) if pPr is not None else None
    if ol is None:
        return
    level = int(ol.get(qn("w:val"))) + 3
    title = "".join(t.text or "" for t in p.findall(".//" + W_T)).replace('"', "'").strip()
    pPr.remove(ol)  # keep it out of the master TOC
    if not title:
        return
    for r in (_run(fldchar="begin"),
              _run(instr=f' TC "{title}" \\f {flag} \\l "{level}" '),
              _run(fldchar="end")):
        p.append(r)


# ── merge ────────────────────────────────────────────────────────────────────

def _set_columns(sectPr, num, space=720):
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        pgmar = sectPr.find(qn("w:pgMar"))
        (pgmar.addnext(cols) if pgmar is not None else sectPr.append(cols))
    cols.set(qn("w:num"), str(num)); cols.set(qn("w:space"), str(space)); cols.set(qn("w:equalWidth"), "1")


def _set_sect_type(sectPr, val):
    t = sectPr.find(qn("w:type"))
    if t is None:
        t = OxmlElement("w:type")
        pgsz = sectPr.find(qn("w:pgSz"))
        (pgsz.addprevious(t) if pgsz is not None else sectPr.insert(0, t))
    t.set(qn("w:val"), val)

def combine(rules_path, comp_path):
    A = Document(rules_path)
    B = Document(comp_path)
    ab, bb = A.element.body, B.element.body

    # Close A's (portrait) section: move its body-level sectPr into a trailing paragraph.
    a_sectPr = ab.findall(W_SECT)[-1]
    p = A.add_paragraph()
    p._p.get_or_add_pPr().append(deepcopy(a_sectPr))
    ab.remove(a_sectPr)

    # Appendix divider at the START of the compendium (the single master-TOC entry),
    # then the compendium body. The compendium's own contents page goes at the very
    # back of the book (below), like an index.
    ab.append(_appendix_heading())
    for child in list(bb):
        c = deepcopy(child)
        if c.tag == W_P:
            _tag_heading_as_tc(c, COMP_FLAG)
        ab.append(c)  # compendium body; its sectPr (if any) stays last
    # Back-of-book index in its OWN 2-column landscape section, TOC levels 3-4 (styled
    # independently of the front TOC). A section break closes the 1-column compendium;
    # the final sectPr (now 2 columns) governs the index, and starts it on a new page.
    tb = _toc_block("Compendium Contents", levels="3-4", flag=COMP_FLAG, title_size=16)
    sect = ab.findall(W_SECT)
    if sect:
        final = sect[-1]                    # body sectPr (landscape, 1 col)
        land = deepcopy(final)
        # closer: ends the 1-column compendium section
        closer = OxmlElement("w:p"); closer.append(OxmlElement("w:pPr"))
        closer.find(W_PPR).append(deepcopy(land))
        # balance: a 2-column section that TERMINATES the index -> Word balances the columns
        # (an unterminated last section fills the first column instead). Starts on a new page.
        bal = deepcopy(land); _set_columns(bal, 2); _set_sect_type(bal, "nextPage")
        balance = OxmlElement("w:p"); balance.append(OxmlElement("w:pPr"))
        balance.find(W_PPR).append(bal)
        # body final sectPr -> empty 1-column trailing section, continuous (no extra page)
        _set_columns(final, 1); _set_sect_type(final, "continuous")
        final.addprevious(closer)
        final.addprevious(tb[1])
        final.addprevious(balance)
    else:
        ab.append(tb[1])

    # Master TOC (whole book) at the very front.
    for idx, el in enumerate(_toc_block("Contents", title_size=18)):
        ab.insert(idx, el)
    return A


def add_footers(doc, version):
    for section in doc.sections:
        section.footer.is_linked_to_previous = False
        p = section.footer.paragraphs[0]
        for r in list(p._p.findall(qn("w:r"))):
            p._p.remove(r)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p._p.append(_run(f"Renown v{version}   \u00b7   Page "))
        p._p.append(_run(fldchar="begin")); p._p.append(_run(instr="PAGE")); p._p.append(_run(fldchar="end"))
        p._p.append(_run(" of "))
        p._p.append(_run(fldchar="begin")); p._p.append(_run(instr="NUMPAGES")); p._p.append(_run(fldchar="end"))


def set_update_fields(doc):
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        uf = OxmlElement("w:updateFields"); uf.set(qn("w:val"), "true"); settings.append(uf)



def style_toc(doc, size=10):
    """Style ONLY the back-index entry styles (TOC 3/4): compact, EB Garamond, with a
    right tab at the column's edge + dot leader so the page number sits next to the
    text in the 2-column index. Front TOC (TOC 1/2) is left as Word's normal styling.
    customStyle is stripped so Word applies these to its generated TOC output."""
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.shared import Inches
    existing = {st.name for st in doc.styles}
    for lvl in (3, 4):
        name = f"TOC {lvl}"
        try:
            st = doc.styles[name] if name in existing else doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except Exception:
            continue
        cs = qn("w:customStyle")
        if cs in st.element.attrib:
            del st.element.attrib[cs]
        pf = st.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
        try:
            st.font.size = Pt(size); st.font.name = FONT
        except Exception:
            pass
        try:
            pf.tab_stops.add_tab_stop(Inches(4.6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        except Exception:
            pass

if __name__ == "__main__":
    if len(sys.argv) < 4:
        sys.exit("usage: python combine_docx.py Rules.docx Compendium.docx Renown.docx [version]")
    version = sys.argv[4] if len(sys.argv) > 4 else VERSION
    doc = combine(sys.argv[1], sys.argv[2])
    add_footers(doc, version)
    set_update_fields(doc)
    style_toc(doc)
    doc.save(sys.argv[3])
    print(f"wrote {sys.argv[3]} (master TOC + compendium TOC via TC marks, footer Renown v{version})")