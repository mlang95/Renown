#!/usr/bin/env python3
"""build_compendium2.py — redesigned Compendium .docx from compendium_data.json.

Improvements over build_compendium.py (all font-independent, verified by render):
  M2  narrow tables paired two-up (reclaims the blank right half of the page)
  M4  split hygiene: rows never split mid-height (cantSplit); header row repeats
      on any table that spans a page break (tblHeader); tables with <=KEEP_MAX
      data rows are glued (keep_with_next) so a small table bumps whole instead of
      orphaning 1-2 rows onto the next page
  Z   zebra striping on long tables (>=ZEBRA_MIN rows) for row tracking
  M1  column sizing: near-empty columns (mostly "—") no longer inherit a wide
      header's width; a small safety factor guards against font-substitution wrap;
      full usable width (10.4") used instead of 10.0"

Usage:  python build_compendium2.py compendium_data.json Compendium.docx
"""
import json, re, sys, os
sys.path.insert(0, os.getcwd())
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
try:
    import renown_data as rd
except Exception as e:
    sys.stderr.write(f"WARN: renown_data not importable ({e}); some tables limited\n")
    rd = None
from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "EB Garamond"
BODY = 8.5          # pt
HEAD_FILL = "EFEFEF"
ZEBRA_FILL = "F4F4F4"
TOTAL = 10.4        # usable width in landscape A4 @ 0.5" margins (10.69" avail)
KEEP_MAX = int(os.environ.get("KEEP_MAX","8"))        # tables with <= this many data rows are glued (never split)
ZEBRA_MIN = 10      # zebra-stripe tables with >= this many data rows

# ── low-level cell/row helpers ───────────────────────────────────────────────
def _set_cell_border(cell, color="auto", sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)

def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def _cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(OxmlElement("w:cantSplit"))

def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:tblHeader")) is None:
        e = OxmlElement("w:tblHeader"); e.set(qn("w:val"), "true"); trPr.append(e)

def _rich(paragraph, text, bold=False, italic=False, size=BODY, keep=False):
    text = str(text)
    parts = [p for p in re.split(r"(\*\*[^*]+\*\*)", text) if p != ""]
    if not parts:
        parts = [text]
    for p in parts:
        m = re.match(r"^\*\*([^*]+)\*\*$", p)
        r = paragraph.add_run(m.group(1) if m else p)
        r.font.name = FONT; r.font.size = Pt(size)
        r.bold = True if m else bold
        r.italic = italic
    if keep:
        paragraph.paragraph_format.keep_with_next = True

def _cell_para(cell, keep=False):
    para = cell.paragraphs[0]
    pf = para.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    if keep:
        pf.keep_with_next = True
    return para

def _alpha(rows, col=0):
    return sorted(rows, key=lambda r: re.sub(r"\*\*","",str(r[col] or "")).lower())

# ── width measurement (real EB Garamond metrics when available) ───────────────
from reportlab.pdfbase.pdfmetrics import stringWidth as _sw
from reportlab.pdfbase import pdfmetrics as _pm
from reportlab.pdfbase.ttfonts import TTFont as _TTF
_MEAS_RF, _MEAS_BF = "Helvetica", "Helvetica-Bold"
for _dir in (os.path.join(os.path.dirname(os.path.abspath(__file__)), "fonts"),
             os.path.dirname(os.path.abspath(__file__)), "fonts", "."):
    try:
        _pm.registerFont(_TTF("EBG", os.path.join(_dir, "EBGaramond-Regular.ttf")))
        _pm.registerFont(_TTF("EBGB", os.path.join(_dir, "EBGaramond-Bold.ttf")))
        _MEAS_RF, _MEAS_BF = "EBG", "EBGB"
        break
    except Exception:
        continue

def _col_layout(headers, rows, total_in=TOTAL, cap=48, tight=None):
    """Fixed column widths that minimize wrapping, measured in the actual font.
    Near-empty columns (mostly em-dash) are sized to their header WORD, not the full
    header, so a wide header can't starve a neighbour. A safety factor guards against
    font substitution. Widest text columns absorb any shrink; short columns keep
    natural width and never wrap mid-word."""
    tight = tight or set()
    FS = 8.5; PAD = 0.20; SF = 1.03
    RF, BF = _MEAS_RF, _MEAS_BF
    def w_full(s, bold): return _sw(re.sub(r"\*\*","",s), BF if bold else RF, FS)*SF/72.0 + PAD
    def w_word(s, bold):
        ws = re.sub(r"\*\*","",s).split()
        return (max((_sw(w, BF if bold else RF, FS) for w in ws), default=0))*SF/72.0 + PAD
    def emptyish(cells):
        vals = [c for c in cells]
        if not vals: return True
        blank = sum(1 for c in vals if re.sub(r"\*\*","",str(c)).strip() in ("", "\u2014", "-", "\u2013"))
        return blank >= 0.8*len(vals)
    n = len(headers)
    nat_in = [0.0]*n; word_in = [0.0]*n
    for i in range(n):
        h = str(headers[i])
        cells = [str(r[i]) for r in rows if i < len(r) and r[i] is not None]
        hnat = w_word(h, True) if emptyish(cells) else w_full(h, True)
        nat_in[i]  = max([hnat] + [w_full(c, False) for c in cells] or [hnat])
        word_in[i] = max([w_word(h, True)] + [w_word(c, False) for c in cells])
    width = [None]*n
    for i in tight: width[i] = nat_in[i]
    fixed = sum(width[i] for i in tight)
    free = [i for i in range(n) if i not in tight]
    rem = total_in - fixed
    if not free:
        s = sum(width); return [w*total_in/s for w in width] if s else width
    if rem <= 0:
        for i in free: width[i] = word_in[i]
        s = sum(width); return [w*total_in/s for w in width]
    if sum(nat_in[i] for i in free) <= rem:
        for i in free: width[i] = nat_in[i]
        return width
    base = sum(word_in[i] for i in free)
    if base >= rem:
        for i in free: width[i] = word_in[i]
        s = fixed + sum(word_in[i] for i in free); k = total_in/s
        return [w*k if i in tight else word_in[i]*k for i, w in enumerate(width)]
    water = rem - base
    lo, hi = 0.0, max(nat_in[i]-word_in[i] for i in free)
    for _ in range(60):
        mid = (lo+hi)/2
        if sum(min(nat_in[i]-word_in[i], mid) for i in free) < water: lo = mid
        else: hi = mid
    for i in free: width[i] = word_in[i] + min(nat_in[i]-word_in[i], lo)
    return width

def _set_col_widths(t, widths_in):
    tbl = t._tbl; tblPr = tbl.tblPr
    lay = tblPr.find(qn("w:tblLayout"))
    if lay is None:
        lay = OxmlElement("w:tblLayout"); tblPr.append(lay)
    lay.set(qn("w:type"), "fixed")
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), str(int(sum(widths_in)*1440))); tblW.set(qn("w:type"), "dxa")
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc in list(grid.findall(qn("w:gridCol"))):
            grid.remove(gc)
        for wi in widths_in:
            gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(int(wi*1440))); grid.append(gc)
    for i, col in enumerate(t.columns):
        wtw = Twips(int(widths_in[i]*1440))
        for cell in col.cells:
            cell.width = wtw

def _fill_table(t, headers, rows, glue, zebra):
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        p = _cell_para(hdr[i], keep=glue); _rich(p, h, bold=True, keep=glue)
        _shade(hdr[i], HEAD_FILL); _set_cell_border(hdr[i])
    _repeat_header(t.rows[0])
    if glue:
        _cant_split(t.rows[0])
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        stripe = zebra and (ri % 2 == 1)
        for i, val in enumerate(row):
            if i >= len(cells): break
            p = _cell_para(cells[i], keep=glue)
            _rich(p, "" if val is None else val, keep=glue)
            _set_cell_border(cells[i])
            if stripe:
                _shade(cells[i], ZEBRA_FILL)
        _cant_split(t.rows[ri+1])

def add_table(doc, headers, rows, total_in=TOTAL, gap=True, tight=None, zebra="auto"):
    if zebra == "auto":
        zebra = len(rows) >= ZEBRA_MIN
    glue = len(rows) <= KEEP_MAX
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.autofit = False
    _fill_table(t, headers, rows, glue, zebra)
    _set_col_widths(t, _col_layout(headers, rows, total_in, tight=tight))
    if gap:
        g = doc.add_paragraph(); g.paragraph_format.space_after = Pt(2)
    return t

def add_tables_row(doc, specs, widths_in=None, gap=True):
    """Place several small tables side by side in one outer row (kept together)."""
    n = len(specs)
    widths_in = widths_in or [TOTAL/n]*n
    outer = doc.add_table(rows=1, cols=n); outer.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = outer._tbl.tblPr
    lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "fixed"); tblPr.append(lay)
    for i, spec in enumerate(specs):
        headers, rows = spec[0], spec[1]
        zebra = spec[2] if len(spec) > 2 else (len(rows) >= ZEBRA_MIN)
        cell = outer.rows[0].cells[i]; cell.width = Twips(int(widths_in[i]*1440))
        cell._tc.remove(cell.paragraphs[0]._p)
        it = cell.add_table(rows=1, cols=len(headers))
        it.alignment = WD_TABLE_ALIGNMENT.LEFT; it.autofit = False
        _fill_table(it, headers, rows, glue=False, zebra=zebra)
        _set_col_widths(it, _col_layout(headers, rows, widths_in[i]-0.14))
    _cant_split(outer.rows[0])
    if gap:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return outer

def _heading(doc, text, size, before, after, level=1, rule=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.keep_with_next = True
    r = p.add_run(text); r.bold = True; r.font.name = FONT; r.font.size = Pt(size)
    ol = OxmlElement("w:outlineLvl"); ol.set(qn("w:val"), str(level - 1))
    p._p.get_or_add_pPr().append(ol)
    if rule:  # thin bottom border under H1 for scannability
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr"); bo = OxmlElement("w:bottom")
        bo.set(qn("w:val"), "single"); bo.set(qn("w:sz"), "6")
        bo.set(qn("w:space"), "2"); bo.set(qn("w:color"), "808080")
        pbdr.append(bo); pPr.append(pbdr)
    return p

def h1(doc, t): return _heading(doc, t, 16, 10, 3, level=1, rule=True)
def h2(doc, t): return _heading(doc, t, 13, 6, 2, level=2)

def build(data, out_path):
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = FONT; style.font.size = Pt(BODY)
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Twips(16838); sec.page_height = Twips(11906)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Twips(720))

    p = doc.add_paragraph(); r = p.add_run("Renown — Compendium")
    r.bold = True; r.font.name = FONT; r.font.size = Pt(22)
    ver = data.get("version", "")
    sub = ("v"+ver+" · " if ver else "") + "Generated from renown_data.py"
    _rich(doc.add_paragraph(), sub, italic=True)

    eq = data["equipment"]
    HALF = [TOTAL/2, TOTAL/2]

    # ── row-prep for each block (rows only; emitters below decide layout) ──
    def eras_spec():
        if rd is not None and getattr(rd, "ERAS", None):
            rows = [[nm, str(e.get("renown","")), str(e.get("armies","")), str(e.get("cities","")),
                     str(e.get("max_settlements","")), f"+{e.get('influence_per_turn',0)}",
                     str(e.get("innate_diplomacy_influence","")), e.get("envoys","") or "",
                     e.get("unlocks","") or "\u2014"] for nm, e in rd.ERAS.items()]
            return ["Era","Renown","Armies","Cities","Max Settlements","Infl/Turn","Diplo Infl","Envoys","Unlocks"], rows
        return ["Era","Renown","Armies","Cities","Infl/Turn","Diplo Infl","Envoys","Unlocks"], data["eras"]

    def retinues_spec():
        if rd is not None and getattr(rd, "RETINUES", None):
            rows = [[n, str(x.get("cost","")), f"{x['to_hit']}+", str(x.get("endurance","")),
                     f"{x['shaking']}+", str(x.get("speed","\u2014")), str(x.get("max_size","\u2014"))]
                    for n, x in rd.RETINUES.items()]
            return ["Retinue","Cost","To Hit","Endurance","Morale","Speed","Max Size"], rows
        return ["Retinue","Cost","To Hit","Endurance","Morale","Speed"], eq["Retinues"]

    def bandits_spec():
        g = rd.BANDIT_GROWTH_PER_ERA; growth = dict(g.items() if isinstance(g, dict) else g)
        equip = getattr(rd, "BANDIT_EQUIPMENT_PER_ERA", {}) or {}
        order = list(rd.ERAS.keys()) if getattr(rd, "ERAS", None) else list(growth.keys())
        rows = [[era, f"+{growth.get(era,'\u2014')}/turn", equip.get(era, "\u2014")] for era in order]
        return ["Era","Bandit Growth","Armaments"], rows

    def domain_standing_spec():
        _emp = data["domain_board"]; _cmb = {r[0]: r for r in data["standing_effects"]}
        merged = []
        for er in _emp:
            dom = er[0]; cr = _cmb.get(dom, [dom,"","",""]); row = [dom]
            for i in (1,2,3):
                e = (er[i] if i < len(er) and er[i] else "").strip()
                c = (cr[i] if i < len(cr) and cr[i] else "").strip()
                row.append((e+" "+c).strip() if c else e)
            merged.append(row)
        return ["Domain","Rising (3)","Established (6)","Sovereign (10)"], merged

    def terrain_spec():
        tac = getattr(rd, "TACTICAL_TERRAIN", {}) or {}
        BASE = {"Hill":"Grassland","Open Field":"Grassland","Mire":"Wetlands","Forest":"Forest",
                "Tundra":"Tundra","Mountains":"Mountains","Water":"Water"}
        battle = {}
        for feat, fd in tac.items():
            base = BASE.get(feat, feat); note = fd.get("effect",""); label = feat if feat != base else ""
            battle.setdefault(base, []).append((f"{label}: {note}" if label else note))
        rows = []
        for n, d in rd.TERRAIN.items():
            mats = ", ".join(d.get("Raw Materials", []) or []) or "\u2014"
            rows.append([n, mats, d.get("Effect","\u2014") or "\u2014", "  ".join(battle.get(n, [])) or "\u2014"])
        return ["Terrain","Raw Materials","Map Effect","Battle Effect"], rows

    # ── assembly ──
    h1(doc, "Progression")
    h2(doc, "Eras"); add_table(doc, *eras_spec())
    h2(doc, "Seasons"); add_table(doc, ["Season","Name","Effect"], data["seasons"])

    h1(doc, "Domains & Standings")
    add_table(doc, *domain_standing_spec())

    h1(doc, "Empire")
    h2(doc, "Settlements"); add_table(doc, ["Settlement","Tier","Sea Variant","Tax","Muster","Build","Wards","Reach","Notes"], data["settlements"])
    h2(doc, "Infrastructure"); add_table(doc, ["Infrastructure","Upkeep","Freq","Empire Bonus","Tier","Build","Requirement"], data["infrastructure"])
    h2(doc, "Wonders"); add_table(doc, ["Wonder","Empire Bonus","Build","Requirement"], data["wonders"])
    if rd is not None and getattr(rd, "TERRAIN", None):
        h2(doc, "Terrain"); add_table(doc, *terrain_spec())

    h1(doc, "Pursuits")
    for s in data["pursuit_sections"]:
        h2(doc, s["title"]); add_table(doc, ["Pursuit","Mastery Unlock","Innate Effect","Mastery Effect"], _alpha(s["rows"]))

    h1(doc, "Economy")
    # Public Order ‖ Faith & Doubt (both narrow) — reclaim right half
    h2(doc, "Public Order  ·  Faith & Doubt Sources")
    add_tables_row(doc, [
        (["PO","State","Effect"], data["public_order"]),
        (["Type","Source","Condition"], data["po_modifiers"]),
    ], widths_in=HALF)
    h2(doc, "Trade & Income"); add_table(doc, ["Rule","Value"], data["trade_rules"], tight={0})

    h1(doc, "Armies & Combat")
    # Retinues ‖ Armor (both short stat blocks) — reclaim right half
    h2(doc, "Retinues  ·  Armor")
    add_tables_row(doc, [
        retinues_spec(),
        (["Armor","Tier","Save","Keywords"], eq["Armor"]),
    ], widths_in=HALF)
    h2(doc, "Melee Weapons"); add_table(doc, ["Weapon","Tier","AP","Init","Keywords"], eq["Weapons"])
    # Ranged ‖ Shields (both short)
    h2(doc, "Ranged Weapons  ·  Shields")
    add_tables_row(doc, [
        (["Ranged","Tier","AP","Init","Keywords"], eq["Ranged"]),
        (["Shield","Tier","Save","Init","Keywords"], eq["Shields"]),
    ], widths_in=HALF)
    h2(doc, "Tactic Matrix")
    def _relabel(v):
        s = str(v); return s.replace("TS","Save").replace("TH","Strike")
    add_table(doc, [_relabel(h) for h in data["tactic_matrix_header"]],
              [[_relabel(c) for c in r] for r in data["tactic_matrix_rows"]])
    lp = doc.add_paragraph(); lp.paragraph_format.space_after = Pt(2)
    _leg = lp.add_run("I = Initiative · Strike = to Strike · Save = to Save"); _leg.italic = True; _leg.font.size = Pt(7)

    if rd is not None and getattr(rd, "BANDIT_GROWTH_PER_ERA", None):
        h1(doc, "Bandits"); add_table(doc, *bandits_spec())

    h1(doc, "Factions")
    add_table(doc, ["Faction","Mechanic"], data["factions"], tight={0})

    h1(doc, "Glossary")
    for cat in data["glossary_categorized"]:
        h2(doc, cat["title"]); add_table(doc, ["Term","Definition"], _alpha(cat["rows"]), tight={0})

    doc.save(out_path)
    print("wrote " + out_path)

if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else "compendium_data.json"
    out = sys.argv[2] if len(sys.argv) > 2 else "Compendium.docx"
    build(json.load(open(src, encoding="utf-8")), out)