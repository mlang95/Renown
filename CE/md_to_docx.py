#!/usr/bin/env python3
"""md_to_docx.py — render a prose Markdown rules doc to a styled .docx, injecting
data tables generated live from renown_data wherever a marker appears.

Single source of truth split:
  - prose          -> the .md (RULES_reorganized.md)
  - data tables    -> renown_data.py, via docx_tables.{{TABLE:name}} / {{GLOSSARY}}
  - inline subs    -> {{DEF:term}}, {{VERSION}}

This replaces the build_docs.py + Rules_authored.docx path: the prose now lives
in the same .md the wiki consumes, so docx and wiki cannot diverge in wording,
and every embedded table regenerates from canon.

Usage:  python md_to_docx.py RULES_reorganized.md Rules.docx
"""
import re, sys
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
import docx_tables as dt
try:
    import renown_data as _rd
    VERSION = str(getattr(_rd, "VERSION", ""))
except Exception:
    VERSION = ""

FONT = "EB Garamond"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
HEAD_SIZES = {1: 20, 2: 18, 3: 16, 4: 14, 5: 12, 6: 11}   # matches reference Rules.docx
HEAD_BEFORE = {1: 12, 2: 12, 3: 8, 4: 6, 5: 6, 6: 6}

TABLE_MK    = re.compile(r"^\s*\{\{TABLE:([a-z_]+)\}\}\s*$")
GLOSSARY_MK = re.compile(r"^\s*\{\{GLOSSARY\}\}\s*$")
ACTIONS_MK  = re.compile(r"^\s*\{\{ACTIONS:([A-Za-z]+)\}\}\s*$")
LIST_MK     = re.compile(r"^\s*\{\{LIST:([A-Z_]+)\}\}\s*$")
COLS_MK     = re.compile(r"^\s*\{\{COLS:(\d)\}\}\s*$")
DEF_MK      = re.compile(r"\{\{DEF:([^}]+)\}\}")
VAL_MK      = re.compile(r"\{\{VAL:([^}]+)\}\}")
VERSION_MK  = re.compile(r"\{\{VERSION\}\}")
INLINE      = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*|`.+?`)")


def _subs(text):
    """Inline {{VERSION}} / {{VAL:path}} / {{DEF:term}} substitutions (run before run parsing)."""
    text = VERSION_MK.sub(VERSION, text)
    text = VAL_MK.sub(lambda m: dt.value(m.group(1).strip()), text)
    text = DEF_MK.sub(lambda m: dt.definition(m.group(1).strip()) or m.group(0), text)
    return text


def _runs(paragraph, text, base_bold=False, base_size=None):
    """Add runs to a paragraph, parsing **bold**, *italic*, ***both***, `code`."""
    text = _subs(text)
    for part in INLINE.split(text):
        if part == "":
            continue
        bold, ital, body = base_bold, False, part
        if part.startswith("***") and part.endswith("***"):
            bold, ital, body = True, True, part[3:-3]
        elif part.startswith("**") and part.endswith("**"):
            bold, body = True, part[2:-2]
        elif part.startswith("*") and part.endswith("*"):
            ital, body = True, part[1:-1]
        elif part.startswith("`") and part.endswith("`"):
            body = part[1:-1]
        r = paragraph.add_run(body)
        r.font.name = FONT
        if base_size:
            r.font.size = Pt(base_size)
        r.bold = bold
        r.italic = ital


def _colbreak(n):
    """A continuous section break. Per OOXML, a sectPr in a paragraph defines the
    section ENDING at that paragraph, so {{COLS:N}} sets N columns for everything
    since the previous break up to here. The final (body) sectPr stays 1-column."""
    cols = ('<w:cols w:space="708"/>' if n == 1
            else f'<w:cols w:num="{n}" w:space="360" w:equalWidth="1"/>')
    return ('<w:p><w:pPr><w:sectPr><w:type w:val="continuous"/>'
            '<w:pgSz w:w="11906" w:h="16838"/>'
            '<w:pgMar w:top="720" w:right="720" w:bottom="720" w:left="720" '
            'w:header="708" w:footer="708" w:gutter="0"/>'
            f'{cols}</w:sectPr></w:pPr></w:p>')


def _inject(doc, ooxml):
    """Insert one or more OOXML sibling elements (a <w:tbl> or several <w:p>) at the
    current end of body content — before the trailing <w:sectPr>, where add_paragraph
    also inserts — so injected tables land at their marker position, not document end."""
    root = parse_xml(f'<w:root xmlns:w="{W_NS}">{ooxml}</w:root>')
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    for child in list(root):
        if sectPr is not None:
            sectPr.addprevious(child)
        else:
            body.append(child)


def _gfm_cells(row):
    row = row.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    return [c.strip() for c in row.split("|")]


def render(md_path, out_path):
    lines = open(md_path, encoding="utf-8").read().split("\n")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Twips(11906); sec.page_height = Twips(16838)   # A4
    sec.top_margin = sec.bottom_margin = Twips(720)                # 0.5"
    sec.left_margin = sec.right_margin = Twips(720)
    sec.header_distance = Twips(708); sec.footer_distance = Twips(708)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0

    i, n = 0, len(lines)
    list_style = {"-": "List Bullet", "*": "List Bullet"}
    while i < n:
        raw = lines[i]
        ln = raw.rstrip()
        s = ln.strip()
        if not s:
            i += 1
            continue

        # block markers
        m = TABLE_MK.match(ln)
        if m:
            _inject(doc, dt.get(m.group(1)))
            i += 1
            continue
        if GLOSSARY_MK.match(ln):
            _inject(doc, dt.glossary_block())
            i += 1
            continue
        am = ACTIONS_MK.match(ln)
        if am:
            _inject(doc, dt.actions(am.group(1)))
            i += 1
            continue
        lm0 = LIST_MK.match(ln)
        if lm0:
            _inject(doc, dt.list_block(lm0.group(1)))
            i += 1
            continue
        cm = COLS_MK.match(ln)
        if cm:
            _inject(doc, _colbreak(int(cm.group(1))))
            i += 1
            continue

        # heading
        hm = re.match(r"^(#{1,6})\s+(.*)$", s)
        if hm:
            lvl = len(hm.group(1))
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(HEAD_BEFORE.get(lvl, 6))
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            _runs(p, hm.group(2), base_bold=True, base_size=HEAD_SIZES.get(lvl, 11))
            ol = OxmlElement("w:outlineLvl"); ol.set(qn("w:val"), str(lvl - 1))
            p._p.get_or_add_pPr().append(ol)
            i += 1
            continue

        # GFM pipe table: header row, separator (|---|), then body rows
        if "|" in s and i + 1 < n and re.match(r"^\s*\|?[\s:\-|]+\|[\s:\-|]*$", lines[i + 1].strip()) and "-" in lines[i + 1]:
            header = _gfm_cells(s)
            i += 2
            body = []
            while i < n and "|" in lines[i] and lines[i].strip():
                body.append(_gfm_cells(lines[i]))
                i += 1
            _inject(doc, dt._table(header, body))   # same dense style as data tables
            continue

        # list item
        lm = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", raw)
        if lm:
            marker = lm.group(2)
            if marker[0].isdigit():
                # Keep the source's own numbering. Word's "List Number" style uses
                # one document-wide counter, so successive lists would read 34,35…
                # instead of restarting at 1; emit the authored number as text.
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
                p.paragraph_format.first_line_indent = Pt(-18)
                _runs(p, f"{marker} {lm.group(3)}")
            else:
                p = doc.add_paragraph(style="List Bullet")
                _runs(p, lm.group(3))
            i += 1
            continue

        # paragraph
        p = doc.add_paragraph()
        _runs(p, s)
        i += 1

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    if len(sys.argv) < 3:
        sys.exit("usage: python md_to_docx.py RULES_reorganized.md Rules.docx")
    out = render(sys.argv[1], sys.argv[2])
    print("wrote " + out)